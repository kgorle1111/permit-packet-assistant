"""Deterministic multi-agency packet assembler. NO LLM in this file — the fill
path reads only the validated canonical record, because a hallucinated parcel
number on a federal form is the one failure this product can never have.

Each agency template maps form labels -> canonical fields. Unfilled or
needs_review fields are rendered as **[CONSULTANT REVIEW: field]** so nothing
leaves the system looking complete when it isn't. Every page is watermarked
DRAFT; the consultant reviews and submits — the tool never files anything.
"""
import json
from datetime import datetime, timezone
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
TEMPLATES = ROOT / "templates" / "agency_forms.json"
PACKETS = ROOT / "packets"

DRAFT_NOTICE = ("DRAFT — assembled by software from client intake; NOT reviewed, NOT submitted. "
                "Requires consultant review before any agency filing.")


def load_templates():
    data = json.loads(TEMPLATES.read_text())
    return {k: v for k, v in data.items() if not k.startswith("_")}


def fill(record: dict, needs_review: list[str], template: dict):
    """Map one agency's form labels onto the canonical record.

    Returns (rows, flags): rows are (label, value) with review markers inline;
    flags lists every label needing consultant attention.
    """
    rows, flags = [], []
    for label, source in template["fields"].items():
        sources = source if isinstance(source, list) else [source]
        vals, flagged = [], False
        for s in sources:
            v = record.get(s)
            if v in (None, "") or s in needs_review:
                flagged = True
                vals.append(f"[CONSULTANT REVIEW: {s}]")
            else:
                vals.append("yes" if v is True else "no" if v is False else str(v))
        rows.append((label, " x ".join(vals) if len(vals) > 1 else vals[0]))
        if flagged:
            flags.append(label)
    return rows, flags


def build_packet(project: dict) -> Path:
    """Assemble the full multi-agency draft packet DOCX for one project."""
    from docx import Document  # deferred import — only pay on export

    record, needs_review = project["fields"], project.get("needs_review", [])
    doc = Document()
    doc.add_heading(f"Permit draft packet — {record.get('owner_name') or project['id']}", level=0)
    doc.add_paragraph(f"Project {project['id']} · Assembled {datetime.now(timezone.utc).strftime('%Y-%m-%d')}")
    doc.add_paragraph(DRAFT_NOTICE, style="Intense Quote")

    total_filled = total_flagged = 0
    for agency, template in load_templates().items():
        rows, flags = fill(record, needs_review, template)
        doc.add_heading(template["title"], level=1)
        t = doc.add_table(rows=0, cols=2)
        for label, value in rows:
            cells = t.add_row().cells
            cells[0].text = label
            cells[1].text = value
        if flags:
            doc.add_paragraph(f"⚠ Consultant review required: {', '.join(flags)}")
        total_filled += len(rows) - len(flags)
        total_flagged += len(flags)

    doc.add_paragraph(f"Auto-filled fields: {total_filled} · Flagged for review: {total_flagged}")
    PACKETS.mkdir(exist_ok=True)
    out = PACKETS / f"{project['id']}.docx"
    doc.save(out)
    return out


def packet_stats(project: dict) -> dict:
    """Fill/flag counts without building the DOCX — feeds the value receipt."""
    record, needs_review = project["fields"], project.get("needs_review", [])
    filled = flagged = 0
    for template in load_templates().values():
        rows, flags = fill(record, needs_review, template)
        filled += len(rows) - len(flags)
        flagged += len(flags)
    return {"fields_autofilled": filled, "fields_flagged": flagged}


if __name__ == "__main__":
    import tempfile
    PACKETS = Path(tempfile.mkdtemp())

    project = {"id": "P1", "needs_review": ["parcel_apn"],
               "fields": {"owner_name": "Pat Lee", "site_address": "12 Cliff Dr, Santa Cruz",
                          "parcel_apn": "007-123-456", "waterbody": "Monterey Bay",
                          "structure_type": "dock", "work_type": "replace",
                          "length_ft": 40.0, "width_ft": 8.0, "existing_structures": "old dock",
                          "setback_notes": None, "riparian_notes": None,
                          "contractor_name": "Bayside Marine", "survey_available": True,
                          "photos_available": None}}
    tpls = load_templates()
    assert set(tpls) == {"county_building", "state_env", "usace_spgp"}
    rows, flags = fill(project["fields"], project["needs_review"], tpls["county_building"])
    as_dict = dict(rows)
    assert as_dict["Assessor parcel number (APN)"].startswith("[CONSULTANT REVIEW")  # needs_review wins over value
    assert "Setback information" in flags                                            # missing -> flagged
    assert as_dict["Structure length (ft)"] == "40.0"
    d2 = dict(fill(project["fields"], [], tpls["state_env"])[0])
    assert d2["Dimensions (L x W, ft)"] == "40.0 x 8.0"
    assert d2["Survey provided"] == "yes"
    path = build_packet(project)
    assert path.exists() and path.stat().st_size > 0
    s = packet_stats(project)
    assert s["fields_autofilled"] > 0 and s["fields_flagged"] > 0
    print("packet OK — flags render inline, nothing fills silently")
